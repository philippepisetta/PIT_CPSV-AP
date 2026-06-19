import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define paths
artifact_dir = r"C:\Users\Philippe Pisetta\.gemini\antigravity\brain\1852d3de-cc1f-4e7e-8d9f-0331e0dfd276"
output_dir = r"c:\Users\Philippe Pisetta\Downloads\testing CPSV-AP"

def parse_markdown_to_docx(md_path, docx_path, title_text):
    doc = Document()
    
    # Set professional styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(title_text)
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 41, 55) # Dark gray #1f2937
    
    doc.add_paragraph() # Spacer
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_headers = []
    table_rows = []
    
    in_code = False
    code_text = ""
    
    for line in lines:
        stripped = line.strip()
        
        # Code block
        if stripped.startswith("```"):
            if in_code:
                # End of code block
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(code_text.strip())
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(75, 85, 99)
                code_text = ""
            else:
                in_code = True
            continue
            
        if in_code:
            code_text += line
            continue
            
        # Table parsing
        if stripped.startswith("|"):
            if "---" in stripped:
                continue # Divider line
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            continue
        elif in_table:
            # End of table
            in_table = False
            # Render table in docx
            if table_headers:
                t = doc.add_table(rows=1 + len(table_rows), cols=len(table_headers))
                t.style = 'Light Shading Accent 1'
                hdr_cells = t.rows[0].cells
                for i, text in enumerate(table_headers):
                    hdr_cells[i].text = text
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                for row_idx, row_data in enumerate(table_rows):
                    row_cells = t.rows[row_idx + 1].cells
                    for i, text in enumerate(row_data):
                        if i < len(row_cells):
                            row_cells[i].text = text
            table_headers = []
            table_rows = []
            
        # Headers
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            # Handle inline bold
            text = stripped[2:]
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
        elif stripped.startswith("1. "):
            p = doc.add_paragraph(style='List Number')
            text = stripped[3:]
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
        elif stripped == "" or stripped == "---":
            continue
        else:
            p = doc.add_paragraph()
            # Handle inline bold and links
            text = line
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Clean markdown link syntax: [text](link) -> text
                    clean_part = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', part)
                    p.add_run(clean_part)
                    
    try:
        doc.save(docx_path)
        print(f"Successfully generated: {docx_path}")
    except PermissionError:
        new_path = docx_path.replace(".docx", "_MAJ.docx")
        doc.save(new_path)
        print(f"File locked. Successfully generated alternative: {new_path}")

# Run generation
parse_markdown_to_docx(
    os.path.join(artifact_dir, "pit_functional_workspaces_guide.md"),
    os.path.join(output_dir, "Documentation_Technique_PIT.docx"),
    "PIT Wallonie - Guide des Workspaces & Spécifications Sémantiques (vNext)"
)

parse_markdown_to_docx(
    os.path.join(artifact_dir, "pit_business_functional_guide.md"),
    os.path.join(output_dir, "Guide_Metier_PIT.docx"),
    "PIT Wallonie - Guide d'Usage Métier & Pilotage Territorial"
)

parse_markdown_to_docx(
    os.path.join(artifact_dir, "catalogue_exhaustif_services.md"),
    os.path.join(output_dir, "Catalogue_Exhaustif_Services_PIT.docx"),
    "PIT Wallonie - Catalogue Exhaustif des Dispositifs Régionaux d'Aide aux Entreprises"
)
